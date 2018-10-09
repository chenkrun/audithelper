#!/usr/bin/env python
# -*- coding: utf-8 -*-

import yaml
import xlrd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from collections import OrderedDict

YAML_CONFIG_PATH = u"配置.yaml"
EXCEL_PATH = u"表格.xlsx"
WORD_PATH = u"文档.docx"
DELIMITER = "-"


class YamlConfig(object):

    def __init__(self, config_path):
        with open(config_path, "r", buffering=-1) as fd:
            self.content_dir = yaml.load(fd)

    def content(self):
        return self.content_dir


class ExcelData(object):

    def __init__(self, excel_path):
        # 如果这里出异常了，例如找不到表格。
        self.work_sheet = xlrd.open_workbook(excel_path).sheet_by_name("work_sheet")

    def tables_content(self):
        '''
            {
                "MF": {
                    "0-0": 现金,
                    "0-1": 123,
                    ... ...,
                },
                ... ...
            }
        '''
        tables_content = OrderedDict()
        for row_num in xrange(self.work_sheet.nrows):
            row_value = self.work_sheet.row_values(row_num)
            tag_index = row_value[0].split(DELIMITER)
            tag = tag_index[0]
            if len(tag_index) == 1:
                tables_content[tag] = {}
                continue
            index = tag_index[1]
            for i in range(1, len(row_value)):
                coordinate = "{}-{}".format(index, str(i-1))
                tables_content[tag][coordinate] = unicode(row_value[i]) #
        return tables_content


class WordTable(object):

    def __init__(self, word_path):
        self.path = word_path
        self.document = Document()
        # 修改word正文字体。
        self.document.styles["Normal"].font.name = u"宋体"
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def create_tables(self, execel_data, config_data):
        for key, value in execel_data.items():
            table_config = config_data.get(key)
            table_row = table_config.get("row")
            table_column = table_config.get("column")
            table = self.document.add_table(rows=table_row, cols=table_column, style="Table Grid")
            for sub_key, sub_value in value.items():
                print sub_key, sub_value
                x_y = sub_key.split(DELIMITER)
                coordinate_x = int(x_y[0])
                coordinate_y = int(x_y[1])
                # add_paragraph函数会另起一个段落，因此会出现一个换行符。
                cell = table.cell(coordinate_x, coordinate_y)
                cell.text = sub_value
                paragraph = cell.paragraphs[0]
                if coordinate_y == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = paragraph.runs[0]
                run.font.size = Pt(table_config.get("content").get("font_size"))

    def save(self):
        self.document.save(self.path)


if __name__ == "__main__":
    excel_data = ExcelData(EXCEL_PATH).tables_content()
    word_data = WordTable(WORD_PATH)
    config_data = YamlConfig(YAML_CONFIG_PATH).content()
    word_data.create_tables(excel_data, config_data)
    word_data.save()
